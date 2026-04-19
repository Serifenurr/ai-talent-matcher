from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Iterable

import docx

_EXPERIENCE_HINTS = (
    "experience",
    "work experience",
    "professional experience",
    "employment",
    "career",
    "iş deneyimi",
    "deneyim",
    "tecrübe",
    "profesyonel deneyim",
    "çalışma geçmişi",
    "kariyer",
)

_OTHER_SECTION_HINTS = (
    "education",
    "academic",
    "qualification",
    "skills",
    "technical skills",
    "projects",
    "certificat",
    "language",
    "references",
    "publication",
    "eğitim",
    "öğrenim",
    "akademik",
    "beceri",
    "yetenek",
    "projeler",
    "sertifika",
    "diller",
    "referans",
    "yayın",
    "özet",
    "summary",
    "profil",
    "hakkımda",
    "about me",
)

_YEAR_RANGE = re.compile(
    r"\b(19\d{2}|20\d{2})\b\s*[-–—/]\s*\b(19\d{2}|20\d{2})\b",
    re.IGNORECASE,
)
_YEAR_TO_PRESENT = re.compile(
    r"\b(19\d{2}|20\d{2})\b\s*[-–—/]\s*"
    r"(present|now|current|ongoing|today|halen|devam|şu\s*an|bugün|günümüz)\b",
    re.IGNORECASE,
)
_PLUS_YEARS = re.compile(
    r"\b(\d{1,2})\s*(\+|plus|yıl|year|sene|years)\b",
    re.IGNORECASE,
)

def read_docx_full_text(file_path: str | Path) -> str:
    doc = docx.Document(str(file_path))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        lines.append(t)
    return "\n".join(lines)

def _line_is_experience_header(line_lower: str) -> bool:
    s = line_lower.strip()
    if len(s) > 80:
        return False
    return any(h in s for h in _EXPERIENCE_HINTS)

def _line_is_other_section_header(line_lower: str) -> bool:
    s = line_lower.strip()
    if len(s) > 80:
        return False
    return any(h in s for h in _OTHER_SECTION_HINTS)

def extract_professional_experience(full_text: str) -> str:
    lines = [ln.strip() for ln in full_text.splitlines() if ln.strip()]
    if not lines:
        return ""

    start: int | None = None
    for i, line in enumerate(lines):
        if _line_is_experience_header(line.lower()):
            start = i + 1
            break

    if start is None:
        return "\n".join(lines)

    end = len(lines)
    for j in range(start, len(lines)):
        if _line_is_other_section_header(lines[j].lower()):
            end = j
            break
    return "\n".join(lines[start:end]).strip()

def compress_professional_text(text: str, max_chars: int = 2200) -> str:
    t = re.sub(r"\s+", " ", text).strip()
    if len(t) <= max_chars:
        return t
    t = t[:max_chars]
    cut = t.rfind(". ")
    if cut > int(max_chars * 0.55):
        return t[: cut + 1].strip()
    return t.strip()

def _parse_year_spans(text: str) -> list[tuple[int, int]]:
    spans: list[tuple[int, int]] = []
    for m in _YEAR_RANGE.finditer(text):
        a, b = int(m.group(1)), int(m.group(2))
        if a > b:
            a, b = b, a
        spans.append((a, b))
    y_now = date.today().year
    for m in _YEAR_TO_PRESENT.finditer(text):
        a = int(m.group(1))
        spans.append((a, y_now))
    return spans

def _merge_spans(spans: Iterable[tuple[int, int]]) -> list[tuple[int, int]]:
    spans = sorted(spans)
    if not spans:
        return []
    out: list[tuple[int, int]] = [spans[0]]
    for s, e in spans[1:]:
        ls, le = out[-1]
        if s <= le + 1:
            out[-1] = (ls, max(le, e))
        else:
            out.append((s, e))
    return out

def estimate_years_experience(text: str) -> float:
    spans = _merge_spans(_parse_year_spans(text))
    total_from_spans = 0.0
    for s, e in spans:
        total_from_spans += max(0.0, float(e - s + 1))
    
    found_years = [total_from_spans]
    for m in _PLUS_YEARS.finditer(text):
        try:
            found_years.append(float(m.group(1)))
        except:
            pass
            
    final_yrs = max(found_years) if found_years else 0.0
    return min(final_yrs, 45.0)

def build_embedding_text(
    doc_head: str,
    experience_compressed: str,
    head_max: int = 450,
    exp_max: int = 2200,
) -> str:
    head = compress_professional_text(doc_head, max_chars=head_max)
    exp = compress_professional_text(experience_compressed, max_chars=exp_max)
    return f"{head}\n\n{exp}".strip()

def cv_isleme_ve_kaydet(cv_yolu: str = "cv_havuzu", index_name: str = "cv-index") -> None:
    import os

    from dotenv import load_dotenv
    from pinecone import Pinecone
    from sentence_transformers import SentenceTransformer

    from metadata_extractor import extract_metadata
    from save_to_cloud import save_to_cloud

    load_dotenv()

    if not os.path.isdir(cv_yolu):
        raise FileNotFoundError(
            f"CV klasörü yok veya bir dizin değil: {cv_yolu!r}. Klasörü oluşturup .docx ekleyin."
        )

    pinecone_key = os.getenv("PINECONE_API_KEY")
    if not pinecone_key:
        raise RuntimeError("PINECONE_API_KEY .env dosyasında tanımlı olmalı.")

    model = SentenceTransformer("all-MiniLM-L6-v2")
    pc = Pinecone(api_key=pinecone_key)
    index = pc.Index(index_name)

    cv_listesi = [f for f in os.listdir(cv_yolu) if f.endswith(".docx")]

    for dosya in cv_listesi:
        dosya_yolu = os.path.join(cv_yolu, dosya)

        metin = read_docx_full_text(dosya_yolu)

        meta = extract_metadata(metin)

        vektor = model.encode(metin).tolist()

        vector_id = dosya.replace(" ", "_").lower().replace(".docx", "")
        onizleme = (meta.get("summary") or "").strip() or compress_professional_text(
            metin, max_chars=500
        )
        skills = meta.get("top_skills") or []
        exp_years = float(int(meta.get("experience_years") or 0))

        md: dict = {
            "filename": dosya,
            "experience_years": exp_years,
            "job_title": (meta.get("job_title") or "")[:200],
            "summary": (meta.get("summary") or "")[:500],
            "text_preview": onizleme[:500]
            if isinstance(onizleme, str)
            else str(onizleme)[:500],
        }
        if skills:
            md["top_skills"] = skills

        save_to_cloud(
            vector_id,
            metin,
            md,
            vector=vektor,
            model=model,
            index=index,
        )

if __name__ == "__main__":
    cv_isleme_ve_kaydet()
